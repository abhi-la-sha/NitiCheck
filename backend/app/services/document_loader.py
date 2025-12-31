"""
Document loading and text extraction service.
Supports PDF and DOCX file formats.
"""
import io
from typing import Optional
from fastapi import UploadFile, HTTPException

try:
    from pypdf import PdfReader
except ImportError:
    # Fallback to PyPDF2 if pypdf not available
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


class DocumentLoader:
    """Handles extraction of text from various document formats."""
    
    @staticmethod
    async def extract_text(file: UploadFile) -> str:
        """
        Extract text from uploaded file.
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If file format is unsupported or extraction fails
        """
        # Read file content
        content = await file.read()
        file_extension = file.filename.split('.')[-1].lower() if file.filename else ''
        
        # Reset file pointer for potential re-reading
        await file.seek(0)
        
        if file_extension == 'pdf':
            return DocumentLoader._extract_from_pdf(content)
        elif file_extension == 'docx':
            return DocumentLoader._extract_from_docx(content)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format. Supported formats: PDF, DOCX"
            )
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file."""
        if PdfReader is None:
            raise HTTPException(
                status_code=500,
                detail="PDF processing library not available"
            )
        
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    # Continue with other pages if one fails
                    continue
            
            if not text_parts:
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract text from PDF. File may be corrupted or image-based."
                )
            
            return "\n\n".join(text_parts)
        
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Error processing PDF: {str(e)}"
            )
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file."""
        if Document is None:
            raise HTTPException(
                status_code=500,
                detail="DOCX processing library not available"
            )
        
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract text from DOCX. File may be empty or corrupted."
                )
            
            return "\n\n".join(text_parts)
        
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=400,
                detail=f"Error processing DOCX: {str(e)}"
            )

